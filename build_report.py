# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x00, 0x72, 0xB2)      # bleu institutionnel
DARK = RGBColor(0x25, 0x23, 0x1F)
GREY = RGBColor(0x5B, 0x56, 0x4D)
GOOD = RGBColor(0x1A, 0xAB, 0x40)
BAD = RGBColor(0xD6, 0x48, 0x3F)
AMBER = RGBColor(0xE8, 0xA3, 0x3D)

doc = Document()

# ---- base style ----
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(6)

for i in range(1, 4):
    hstyle = doc.styles[f'Heading {i}']
    hstyle.font.name = 'Segoe UI Semibold'
    hstyle.font.color.rgb = PRIMARY
    hstyle.font.size = Pt([18, 14, 12][i-1])
    hstyle.font.bold = True
    hstyle.paragraph_format.space_before = Pt([18, 14, 10][i-1])
    hstyle.paragraph_format.space_after = Pt(6)

sections = doc.sections
for s in sections:
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)


def add_shaded_paragraph(text, bg_hex, color=DARK, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)
    return p


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E4DFD5')
    pBdr.append(bottom)
    pPr.append(pBdr)


def kpi_row(table, label, value, note=""):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    row[2].text = note
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)
    row[1].paragraphs[0].runs[0].font.bold = True
    row[1].paragraphs[0].runs[0].font.color.rgb = PRIMARY


# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT D'ANALYSE")
run.font.size = Pt(14)
run.font.color.rgb = GREY
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Performance Programmatique, Financière,\nOpérationnelle et de Redevabilité")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Faso Espoir International (FEI) — Burkina Faso")
run.font.size = Pt(15)
run.font.color.rgb = DARK

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Période analysée : 2021 – 2025  |  Devise de référence : USD")
run.font.size = Pt(11)
run.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rapport destiné à la Direction — 18 juillet 2026")
run.font.size = Pt(11)
run.font.color.rgb = GREY

for _ in range(2):
    doc.add_paragraph()

add_shaded_paragraph(
    "CONFIDENTIEL — Usage interne. Les données nominatives des bénéficiaires ont été exclues de cette analyse ; "
    "seules des données agrégées et anonymisées sont présentées.",
    "F4F1EA", color=BAD, bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_page_break()

# ============================================================
# 1. CONTEXTE
# ============================================================
doc.add_heading("1. Contexte", level=1)
doc.add_paragraph(
    "Faso Espoir International (FEI) est une organisation non gouvernementale internationale (ONGI) intervenant "
    "au Burkina Faso dans le domaine du Suivi, de l'Évaluation, de la Redevabilité et de l'Apprentissage (MEAL). "
    "Depuis 2021, FEI opère à travers 12 bases terrain réparties sur l'ensemble du territoire et couvre 11 secteurs "
    "d'intervention humanitaire (sécurité alimentaire, nutrition, eau-hygiène-assainissement, protection, moyens de "
    "subsistance, entre autres), en coordination avec le cluster OCHA Burkina Faso."
)
doc.add_paragraph(
    "Ce rapport a pour objet de présenter à la Direction une analyse consolidée de la base de données MEAL de "
    "l'organisation, afin d'éclairer la prise de décision sur quatre axes prioritaires : la performance "
    "programmatique, le suivi financier, les opérations et la chaîne d'approvisionnement, ainsi que la redevabilité "
    "et la qualité du programme. Il s'appuie directement sur le tableau de bord Power BI développé pour le suivi "
    "continu de ces indicateurs."
)

# ============================================================
# 2. METHODOLOGIE
# ============================================================
doc.add_heading("2. Méthodologie", level=1)
doc.add_paragraph(
    "L'analyse repose sur la base de données MEAL consolidée de FEI (25 tables couvrant les bénéficiaires, "
    "activités, décaissements, cadre logique, plaintes, stocks, approvisionnements et partenaires de mise en "
    "œuvre), collectée via KoboToolbox/ODK sur la période fiscale 2021-2025. Les données ont fait l'objet d'un "
    "nettoyage et d'une modélisation (correction des types de données, consolidation des relations entre tables, "
    "mise en place d'un calendrier d'analyse temporelle) avant le calcul de 31 indicateurs clés répartis en quatre "
    "familles. Les champs nominatifs des bénéficiaires (nom, prénom, téléphone) ont été exclus de l'analyse ; "
    "seules des catégories agrégées (âge, sexe, région, vulnérabilité) sont mobilisées. Les figures présentées sont "
    "des agrégats calculés sur l'ensemble de la période sauf mention contraire."
)

doc.add_page_break()

# ============================================================
# 3. PRINCIPAUX CONSTATS
# ============================================================
doc.add_heading("3. Principaux constats", level=1)
doc.add_paragraph(
    "Vue d'ensemble des indicateurs clés sur la période 2021-2025 :"
)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Cm(7.5), Cm(3.5), Cm(5.5)]
hdr = table.rows[0].cells
hdr[0].text = "Indicateur"
hdr[1].text = "Valeur"
hdr[2].text = "Lecture"
for i, cell in enumerate(hdr):
    cell.width = widths[i]
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '0072B2')
    cell._tc.get_or_add_tcPr().append(shd)

kpis = [
    ("Bénéficiaires actifs", "1 372 / 2 000 (68,6 %)", "373 inactifs, 255 sortis du programme"),
    ("Taux d'exécution budgétaire", "91,7 %", "24,11 M USD décaissés / 26,29 M USD prévus"),
    ("Taux de réalisation cadre logique (volume)", "70,6 %", "3,55 M bénéficiaires atteints / 5,03 M ciblés"),
    ("Objectifs individuels « Atteint »", "20,6 % (103 / 500)", "42,8 % « Non Atteint »"),
    ("Activités réalisées", "82,2 % (12 325 / 15 000)", "8,1 % reportées, 5,7 % annulées"),
    ("Livraisons en retard", "12,1 %", "délai moyen 3,3 jours"),
    ("Bases en rupture de stock imminente", "11 / 12", "47 mouvements de stock critiques"),
    ("Taux de résolution des plaintes", "56,7 % (1 133 / 2 000)", "19,8 % encore en cours de traitement"),
    ("Score de performance des partenaires", "0,80 / 1", "délai moyen de rapportage : 18,1 jours"),
]
for label, val, note in kpis:
    kpi_row(table, label, val, note)
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

doc.add_paragraph()

# --- 3.1 Programme ---
doc.add_heading("3.1 Performance programmatique", level=2)
doc.add_paragraph(
    "FEI a enregistré 2 000 bénéficiaires depuis 2021, dont 1 372 actifs (68,6 %), 373 inactifs et 255 sortis du "
    "programme. La couverture est concentrée sur trois régions à forte insécurité — Est, Sahel et Centre-Nord — "
    "qui représentent à elles seules 42 % des bénéficiaires enregistrés. 7,4 % des bénéficiaires vivent avec un "
    "handicap et le score de vulnérabilité moyen s'établit à 51/100."
)
doc.add_paragraph(
    "Le cadre logique affiche un taux de réalisation global de 70,6 % en volume (3,55 millions de bénéficiaires "
    "atteints sur 5,03 millions ciblés), ce qui masque une réalité plus contrastée : sur les 500 lignes d'objectifs "
    "suivies individuellement, seules 20,6 % sont marquées « Atteint », 36,6 % « Partiellement Atteint » et "
    "42,8 % « Non Atteint ». L'agrégat global de réalisation est donc porté par un nombre restreint d'objectifs "
    "très performants, tandis qu'une part importante des cibles individuelles sous-performe significativement."
)

# --- 3.2 Finance ---
doc.add_heading("3.2 Suivi financier", level=2)
doc.add_paragraph(
    "Le taux d'exécution budgétaire s'établit à 91,7 % sur la période (24,11 M USD décaissés pour 26,29 M USD "
    "prévus), un niveau solide et stable. Les frais de gestion représentent en moyenne 20 % de la valeur de "
    "l'assistance chaque année depuis 2021 (20,2 % en 2021, 20,05 % en 2025), démontrant une bonne maîtrise des "
    "coûts de structure sur la durée."
)
doc.add_paragraph(
    "Un point de vigilance concerne le circuit de décaissement : sur 15 000 décaissements enregistrés, 1 070 "
    "restent « En Attente » et 470 ont été « Rejetés », soit 10,3 % du volume total qui mérite une revue des "
    "causes de blocage ou de rejet."
)

# --- 3.3 Operations ---
doc.add_heading("3.3 Opérations & chaîne d'approvisionnement", level=2)
doc.add_paragraph(
    "Sur 15 000 activités enregistrées, 82,2 % ont été réalisées, 8,1 % reportées et 5,7 % annulées. Le délai "
    "moyen de livraison est de 3,3 jours, avec 12,1 % des livraisons en retard — un niveau globalement maîtrisé "
    "mais perfectible."
)
doc.add_paragraph(
    "Le signal le plus préoccupant concerne les stocks : 47 mouvements ont été classés en « Rupture Imminente », "
    "touchant 11 des 12 bases terrain — un risque de rupture quasi généralisé sur le réseau logistique. Le taux de "
    "couverture moyen des besoins par intervention est de 70 %, pour une satisfaction moyenne des bénéficiaires de "
    "74,9/100."
)

# --- 3.4 Redevabilite ---
doc.add_heading("3.4 Redevabilité & qualité", level=2)
doc.add_paragraph(
    "2 000 plaintes ont été enregistrées sur la période, dont 56,7 % résolues, 19,8 % encore en cours de "
    "traitement et 8,1 % escaladées. Les trois motifs les plus fréquents sont le ciblage/éligibilité contestée "
    "(218 cas), la discrimination perçue dans le ciblage (210 cas) et les retards de distribution (209 cas) — un "
    "signal fort sur la perception de l'équité du ciblage par les communautés. 56,7 % des plaintes ont donné lieu "
    "à une mesure corrective, et 36 657 USD de compensations ont été versées."
)
doc.add_paragraph(
    "Les 60 partenaires de mise en œuvre affichent un score de performance moyen de 0,80 sur 1, avec un délai "
    "moyen de rapportage de 18,1 jours ; 55 partenaires sur 60 sont actuellement actifs."
)

doc.add_page_break()

# ============================================================
# 4. RECOMMANDATIONS
# ============================================================
doc.add_heading("4. Recommandations", level=1)
doc.add_paragraph(
    "Les recommandations ci-dessous sont priorisées en fonction de l'urgence et de l'impact attendu sur la "
    "performance du programme."
)

def reco_block(priority, color, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{priority}  ")
    run.font.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(9)
    run2 = p.add_run(title)
    run2.font.bold = True
    run2.font.color.rgb = DARK
    run2.font.size = Pt(11)
    body = doc.add_paragraph(text)
    body.paragraph_format.space_after = Pt(12)

reco_block("PRIORITAIRE", BAD, "1. Auditer l'équité et la transparence du ciblage",
    "21,8 % des plaintes portent sur l'éligibilité/le ciblage contesté et 10,5 % sur une discrimination perçue "
    "dans le ciblage. Recommandation : conduire un audit des critères de sélection par base, renforcer la "
    "communication communautaire sur les critères d'éligibilité, et mettre en place un retour systématique aux "
    "plaignants sur les décisions de ciblage.")

reco_block("PRIORITAIRE", BAD, "2. Sécuriser la chaîne d'approvisionnement",
    "11 bases sur 12 ont connu au moins un mouvement de stock en rupture imminente. Recommandation : réviser les "
    "seuils de stock minimum/maximum par intervention, raccourcir les délais de commande auprès des partenaires "
    "d'approvisionnement et mettre en place une alerte automatisée sur le tableau de bord avant que le seuil "
    "critique ne soit atteint.")

reco_block("PRIORITAIRE", BAD, "3. Cibler le pilotage du cadre logique au niveau des objectifs individuels",
    "Le taux de réalisation global (70,6 %) masque le fait que 42,8 % des lignes d'objectifs individuelles sont "
    "« Non Atteint ». Recommandation : identifier les bases et secteurs concentrant la sous-performance pour une "
    "réallocation ciblée des ressources, plutôt que de piloter uniquement sur l'indicateur agrégé.")

reco_block("MOYEN TERME", AMBER, "4. Fiabiliser le circuit de décaissement",
    "10,3 % des décaissements restent en attente ou ont été rejetés. Recommandation : analyser les causes de "
    "rejet par bailleur/base et réduire les délais de validation financière.")

reco_block("MOYEN TERME", AMBER, "5. Accompagner les partenaires de mise en œuvre les moins performants",
    "Le score de performance moyen des partenaires est de 0,80/1 avec un délai moyen de rapportage de 18,1 jours. "
    "Recommandation : mettre en place un plan d'accompagnement pour les partenaires sous le score moyen et "
    "clarifier les délais de rapportage attendus dans les accords de partenariat.")

reco_block("MOYEN TERME", AMBER, "6. Renforcer la couverture des besoins dans les secteurs les plus sollicités",
    "Le taux de couverture moyen des besoins est de 70 % ; les secteurs Nutrition, Moyens de Subsistance et EHA "
    "concentrent le plus grand nombre d'interventions. Recommandation : prioriser le renforcement de capacité "
    "dans ces trois secteurs lors du prochain cycle de programmation.")

reco_block("À MAINTENIR", GOOD, "7. Poursuivre la discipline budgétaire actuelle",
    "Le taux d'exécution budgétaire (91,7 %) et les frais de gestion stables (~20 % depuis 2021) témoignent d'une "
    "bonne gestion financière. Recommandation : documenter cette performance dans les rapports bailleurs comme "
    "preuve de fiabilité de gestion.")

reco_block("À MAINTENIR", GOOD, "8. Maintenir un mécanisme de plaintes actif et accessible",
    "56,7 % des plaintes ont été résolues avec mesures correctives. Recommandation : maintenir l'accessibilité du "
    "mécanisme (Ligne Verte, boîtes à suggestions, comités communautaires) tout en renforçant les capacités de "
    "traitement pour réduire le taux de plaintes « en cours » (19,8 %).")

hr()
p = doc.add_paragraph()
run = p.add_run(
    "Source : Base de données MEAL Faso Espoir International, 2021-2025. Tableau de bord Power BI — "
    "Base_MEAL.pbip. Document confidentiel à usage interne."
)
run.font.size = Pt(8.5)
run.font.color.rgb = GREY
run.font.italic = True

doc.save("Rapport_Analyse_FEI_2021-2025.docx")
print("OK")
